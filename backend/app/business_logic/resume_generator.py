import json
import logging
import re
from pathlib import Path
from typing import Dict, Any, Optional

from jinja2 import Environment, FileSystemLoader
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.database.models import Resume, ResumeVersion, User
from app.schemas.resume import ResumeContentSchema

logger = logging.getLogger(__name__)

TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "templates"
TEMPLATE_NAME = "resume_template.html"


def format_display_url(url: Optional[str]) -> str:
    """Format a web/social URL for clean presentation on resumes (e.g. linkedin.com/in/john)."""
    if not url:
        return ""
    u = str(url).strip()
    u = re.sub(r"^https?://", "", u, flags=re.IGNORECASE)
    u = re.sub(r"^www\.", "", u, flags=re.IGNORECASE)
    return u.rstrip("/")


# ---------------------------------------------------------------------------
# Dynamic spacing engine
# ---------------------------------------------------------------------------
SCALE_TIERS = {
    "spacious": dict(
        fs_base="10.5pt",
        fs_name="22pt",
        fs_title="11pt",
        fs_section="11.5pt",
        fs_sub="9.5pt",
        line_height=1.46,
        section_gap="18px",
        item_gap="12px",
        para_gap="5px",
        page_margin="10mm 12mm",
    ),
    "normal": dict(
        fs_base="10pt",
        fs_name="20pt",
        fs_title="10.5pt",
        fs_section="11pt",
        fs_sub="9pt",
        line_height=1.38,
        section_gap="14px",
        item_gap="9px",
        para_gap="4px",
        page_margin="8mm 10mm",
    ),
    "compact": dict(
        fs_base="9.5pt",
        fs_name="18pt",
        fs_title="10pt",
        fs_section="10.5pt",
        fs_sub="8.5pt",
        line_height=1.28,
        section_gap="10px",
        item_gap="7px",
        para_gap="3px",
        page_margin="7mm 9mm",
    ),
    "dense": dict(
        fs_base="9pt",
        fs_name="17pt",
        fs_title="9.5pt",
        fs_section="10pt",
        fs_sub="8pt",
        line_height=1.20,
        section_gap="8px",
        item_gap="5px",
        para_gap="2px",
        page_margin="6mm 8mm",
    ),
}


def _text_len(value) -> int:
    return len(value) if isinstance(value, str) else 0


def estimate_content_score(data: dict) -> int:
    """Rough weighted count of 'how much stuff is on this resume'."""
    score = 0
    pi = data.get("personal_info") or {}
    score += _text_len(pi.get("summary"))

    for exp in data.get("experience") or []:
        score += 50  # fixed cost per role (title/company/dates lines)
        score += _text_len(exp.get("description"))
        for a in exp.get("achievements") or []:
            score += _text_len(a) + 8

    for edu in data.get("education") or []:
        score += 40
        score += _text_len(edu.get("description"))

    score += len(data.get("skills") or []) * 8

    for proj in data.get("projects") or []:
        score += 40
        score += _text_len(proj.get("description"))
        for a in proj.get("achievements") or []:
            score += _text_len(a) + 8
        score += len(proj.get("technologies") or []) * 5

    for cert in data.get("certifications") or []:
        score += 25

    return score


def pick_scale(data: dict) -> dict:
    score = estimate_content_score(data)
    if score < 1200:
        tier = "spacious"
    elif score < 2200:
        tier = "normal"
    elif score < 3200:
        tier = "compact"
    else:
        tier = "dense"
    return SCALE_TIERS[tier]


def _safe_defaults(data: dict) -> dict:
    """Fill in missing top-level keys so the template never KeyErrors."""
    data = dict(data)
    data.setdefault("personal_info", {})
    data.setdefault("experience", [])
    data.setdefault("education", [])
    data.setdefault("skills", [])
    data.setdefault("projects", [])
    data.setdefault("certifications", [])
    return data


def render_resume_html(data: dict) -> str:
    """Render resume data into dynamic ATS-safe HTML using Jinja2 and content-scaled CSS."""
    data = _safe_defaults(data)
    scale = pick_scale(data)
    
    # Separate technical from soft skills
    from app.business_logic.cv_extractor import categorize_skills
    all_skills = data["skills"]
    technical_skills, soft_skills = categorize_skills(all_skills)

    env = Environment(
        loader=FileSystemLoader(str(TEMPLATE_DIR)),
        autoescape=True,
        trim_blocks=True,
        lstrip_blocks=True,
    )
    env.filters["format_display_url"] = format_display_url
    template = env.get_template(TEMPLATE_NAME)

    return template.render(
        personal_info=data["personal_info"],
        experience=data["experience"],
        education=data["education"],
        technical_skills=technical_skills,
        soft_skills=soft_skills,
        skills=all_skills,  # Keep for backward compatibility
        projects=data["projects"],
        certifications=data["certifications"],
        scale=scale,
        page_margin=scale["page_margin"],
    )


def render_resume_pdf(data: dict, output_path: str) -> None:
    """Render resume data to PDF using WeasyPrint (if installed)."""
    try:
        from weasyprint import HTML
        html_str = render_resume_html(data)
        HTML(string=html_str, base_url=str(TEMPLATE_DIR)).write_pdf(output_path)
    except ImportError:
        logger.warning("WeasyPrint is not installed. PDF generation via server unavailable.")
        raise RuntimeError("WeasyPrint is not installed on this server.")


def render_resume_docx(data: dict) -> "io.BytesIO":
    """Render structured resume data into an ATS-safe Word (.docx) document."""
    import io
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml

    doc = docx.Document()

    # 0.5 inch page margins (ATS standard)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    personal_info = data.get("personal_info") or {}

    def format_run(run, font_name="Calibri", size_pt=10, bold=False, italic=False, color_rgb=(15, 23, 42)):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(*color_rgb)

    # 1. Header (Name, Title, Contact)
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_before = Pt(0)
    name_p.paragraph_format.space_after = Pt(2)
    r = name_p.add_run(personal_info.get("full_name") or "Candidate Name")
    format_run(r, font_name="Calibri", size_pt=20, bold=True, color_rgb=(15, 23, 42))

    if personal_info.get("professional_title"):
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(0)
        title_p.paragraph_format.space_after = Pt(4)
        r = title_p.add_run(personal_info["professional_title"])
        format_run(r, font_name="Calibri", size_pt=11, color_rgb=(51, 65, 85))

    contact_parts = []
    if personal_info.get("phone"): contact_parts.append(personal_info["phone"])
    if personal_info.get("email"): contact_parts.append(personal_info["email"])
    if personal_info.get("location"): contact_parts.append(personal_info["location"])
    if personal_info.get("linkedin_url"): contact_parts.append(format_display_url(personal_info["linkedin_url"]))
    if personal_info.get("github_url"): contact_parts.append(format_display_url(personal_info["github_url"]))
    if personal_info.get("portfolio_url"): contact_parts.append(format_display_url(personal_info["portfolio_url"]))

    if contact_parts:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_before = Pt(0)
        contact_p.paragraph_format.space_after = Pt(8)
        contact_str = "  •  ".join(contact_parts)
        r = contact_p.add_run(contact_str)
        format_run(r, font_name="Calibri", size_pt=9.5, color_rgb=(100, 116, 139))

    def add_section_heading(title: str):
        sec_p = doc.add_paragraph()
        sec_p.paragraph_format.space_before = Pt(10)
        sec_p.paragraph_format.space_after = Pt(4)
        r = sec_p.add_run(title.upper())
        format_run(r, font_name="Calibri", size_pt=10.5, bold=True, color_rgb=(15, 23, 42))
        try:
            pPr = sec_p._p.get_or_add_pPr()
            pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="8" w:space="2" w:color="0F172A"/></w:pBdr>')
            pPr.append(pBdr)
        except Exception:
            pass

    # 2. Summary
    if personal_info.get("summary"):
        add_section_heading("Summary")
        sum_p = doc.add_paragraph()
        sum_p.paragraph_format.space_before = Pt(2)
        sum_p.paragraph_format.space_after = Pt(6)
        r = sum_p.add_run(personal_info["summary"])
        format_run(r, font_name="Calibri", size_pt=10, color_rgb=(51, 65, 85))

    # 3. Experience
    experience = data.get("experience") or []
    if experience:
        add_section_heading("Experience")
        for exp in experience:
            head_p = doc.add_paragraph()
            head_p.paragraph_format.space_before = Pt(4)
            head_p.paragraph_format.space_after = Pt(1)
            r_pos = head_p.add_run(exp.get("position") or "")
            format_run(r_pos, font_name="Calibri", size_pt=10.5, bold=True, color_rgb=(15, 23, 42))

            dates_str = f" ({exp.get('start_date', '')} — {'Present' if exp.get('is_current') else exp.get('end_date', '')})"
            r_date = head_p.add_run(dates_str)
            format_run(r_date, font_name="Calibri", size_pt=9.5, italic=True, color_rgb=(100, 116, 139))

            if exp.get("company") or exp.get("location"):
                sub_p = doc.add_paragraph()
                sub_p.paragraph_format.space_before = Pt(0)
                sub_p.paragraph_format.space_after = Pt(2)
                sub_text = exp.get("company", "") + (f", {exp['location']}" if exp.get("location") else "")
                r_sub = sub_p.add_run(sub_text)
                format_run(r_sub, font_name="Calibri", size_pt=9.5, italic=True, color_rgb=(51, 65, 85))

            for bullet in exp.get("achievements") or []:
                if bullet and bullet.strip() and bullet.strip() not in ("•", "-"):
                    b_p = doc.add_paragraph(style='List Bullet')
                    b_p.paragraph_format.space_before = Pt(1)
                    b_p.paragraph_format.space_after = Pt(2)
                    r_b = b_p.add_run(bullet.strip())
                    format_run(r_b, font_name="Calibri", size_pt=9.5, color_rgb=(51, 65, 85))

    # 4. Projects
    projects = data.get("projects") or []
    if projects:
        add_section_heading("Projects")
        for proj in projects:
            p_p = doc.add_paragraph()
            p_p.paragraph_format.space_before = Pt(4)
            p_p.paragraph_format.space_after = Pt(2)
            r_pn = p_p.add_run(proj.get("name", ""))
            format_run(r_pn, font_name="Calibri", size_pt=10.5, bold=True, color_rgb=(15, 23, 42))

            if proj.get("description"):
                desc_p = doc.add_paragraph()
                desc_p.paragraph_format.space_before = Pt(0)
                desc_p.paragraph_format.space_after = Pt(4)
                r_pd = desc_p.add_run(proj["description"])
                format_run(r_pd, font_name="Calibri", size_pt=9.5, color_rgb=(51, 65, 85))

            for bullet in proj.get("achievements") or []:
                if bullet and bullet.strip() and bullet.strip() not in ("•", "-"):
                    b_p = doc.add_paragraph(style='List Bullet')
                    b_p.paragraph_format.space_before = Pt(1)
                    b_p.paragraph_format.space_after = Pt(2)
                    r_b = b_p.add_run(bullet.strip())
                    format_run(r_b, font_name="Calibri", size_pt=9.5, color_rgb=(51, 65, 85))

    # 5. Skills
    all_skills = data.get("skills") or []
    if all_skills:
        add_section_heading("Skills")
        
        # Separate technical from soft skills
        from app.business_logic.cv_extractor import categorize_skills
        technical_skills, soft_skills = categorize_skills(all_skills)
        
        # Display technical skills
        if technical_skills:
            sk_p = doc.add_paragraph()
            sk_p.paragraph_format.space_before = Pt(2)
            sk_p.paragraph_format.space_after = Pt(2)
            r_pre = sk_p.add_run("Technical Proficiencies: ")
            format_run(r_pre, font_name="Calibri", size_pt=9.5, bold=True, color_rgb=(15, 23, 42))
            r_sk = sk_p.add_run(", ".join(technical_skills))
            format_run(r_sk, font_name="Calibri", size_pt=9.5, color_rgb=(51, 65, 85))
        
        # Display soft skills
        if soft_skills:
            soft_p = doc.add_paragraph()
            soft_p.paragraph_format.space_before = Pt(1)
            soft_p.paragraph_format.space_after = Pt(6)
            r_soft_pre = soft_p.add_run("Professional Skills: ")
            format_run(r_soft_pre, font_name="Calibri", size_pt=9.5, bold=True, color_rgb=(15, 23, 42))
            r_soft = soft_p.add_run(", ".join(soft_skills))
            format_run(r_soft, font_name="Calibri", size_pt=9.5, color_rgb=(51, 65, 85))

    # 6. Education
    education = data.get("education") or []
    if education:
        add_section_heading("Education")
        for edu in education:
            edu_p = doc.add_paragraph()
            edu_p.paragraph_format.space_before = Pt(3)
            edu_p.paragraph_format.space_after = Pt(3)
            edu_title = edu.get("degree", "") + (f" in {edu['field_of_study']}" if edu.get("field_of_study") else "") + (f", {edu['institution']}" if edu.get("institution") else "")
            r_ed = edu_p.add_run(edu_title)
            format_run(r_ed, font_name="Calibri", size_pt=10, bold=True, color_rgb=(15, 23, 42))

            dates_str = f" ({edu.get('start_date', '')} — {'Present' if edu.get('is_current') else edu.get('end_date', '')})"
            r_dt = edu_p.add_run(dates_str)
            format_run(r_dt, font_name="Calibri", size_pt=9.5, italic=True, color_rgb=(100, 116, 139))

    # 7. Certifications
    certifications = data.get("certifications") or []
    if certifications:
        add_section_heading("Certifications")
        for cert in certifications:
            cert_p = doc.add_paragraph()
            cert_p.paragraph_format.space_before = Pt(2)
            cert_p.paragraph_format.space_after = Pt(2)
            r_c = cert_p.add_run(cert.get("name", ""))
            format_run(r_c, font_name="Calibri", size_pt=9.5, bold=True, color_rgb=(15, 23, 42))
            if cert.get("issuing_organization"):
                r_o = cert_p.add_run(f" — {cert['issuing_organization']}")
                format_run(r_o, font_name="Calibri", size_pt=9.5, color_rgb=(51, 65, 85))
            if cert.get("issue_date"):
                r_cd = cert_p.add_run(f" ({cert['issue_date']})")
                format_run(r_cd, font_name="Calibri", size_pt=9, italic=True, color_rgb=(100, 116, 139))

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def prepare_resume_render_data(profile_snapshot: Dict[str, Any], content: Dict[str, Any]) -> Dict[str, Any]:
    """Merge profile snapshot and edited resume content into standard structure for HTML/DOCX rendering."""
    personal_info = dict(content.get("personal_info") or profile_snapshot.get("personal_info") or {})
    if not personal_info.get("summary") and content.get("summary"):
        personal_info["summary"] = content.get("summary")

    return {
        "personal_info": personal_info,
        "experience": content.get("experience") if content.get("experience") is not None else profile_snapshot.get("experience", []),
        "education": content.get("education") if content.get("education") is not None else profile_snapshot.get("education", []),
        "skills": content.get("skills") if content.get("skills") is not None else profile_snapshot.get("skills", []),
        "projects": content.get("projects") if content.get("projects") is not None else profile_snapshot.get("projects", []),
        "certifications": content.get("certifications") if content.get("certifications") is not None else profile_snapshot.get("certifications", []),
    }


# ---------------------------------------------------------------------------
# Direct Resume Content Generator (Pure Data -> Schema, No LLM Call Needed)
# ---------------------------------------------------------------------------
async def generate_ats_resume_content(profile_dict: Dict[str, Any], target_role: str = "", custom_instructions: str = "") -> ResumeContentSchema:
    """Generate ATS-optimized resume content directly from candidate profile data without extra LLM call."""
    personal_info = profile_dict.get("personal_info", {})
    name = personal_info.get("full_name", "Candidate")
    title = target_role or personal_info.get("professional_title", "Professional")
    existing_summary = personal_info.get("summary", "")

    summary = existing_summary if len(existing_summary) > 10 else f"Results-driven {title} with expertise in building scalable solutions and delivering high-impact results."

    experiences = []
    for exp in profile_dict.get("experience", []):
        achievements = exp.get("achievements", [])
        if not achievements and exp.get("description"):
            achievements = [exp["description"]]

        experiences.append({
            "company": exp.get("company", ""),
            "position": exp.get("position", ""),
            "location": exp.get("location", ""),
            "employment_type": exp.get("employment_type", ""),
            "start_date": exp.get("start_date", ""),
            "end_date": exp.get("end_date", "Present"),
            "is_current": exp.get("is_current", False),
            "description": exp.get("description", ""),
            "achievements": achievements or ["Executed core responsibilities and achieved team milestones."]
        })

    education = []
    for edu in profile_dict.get("education", []):
        education.append({
            "institution": edu.get("institution", ""),
            "degree": edu.get("degree", ""),
            "field_of_study": edu.get("field_of_study", ""),
            "start_date": edu.get("start_date", ""),
            "end_date": edu.get("end_date", "Present"),
            "is_current": edu.get("is_current", False),
            "gpa": edu.get("gpa", ""),
            "description": edu.get("description", "")
        })

    projects = []
    for proj in profile_dict.get("projects", []):
        achievements = proj.get("achievements", [])
        if not achievements and proj.get("description"):
            achievements = [proj["description"]]
        projects.append({
            "name": proj.get("name", ""),
            "technologies": proj.get("technologies", []),
            "description": proj.get("description", ""),
            "project_url": proj.get("project_url", ""),
            "github_url": proj.get("github_url", ""),
            "achievements": achievements or []
        })

    certifications = []
    for cert in profile_dict.get("certifications", []):
        certifications.append({
            "name": cert.get("name", ""),
            "issuing_organization": cert.get("issuing_organization", ""),
            "issue_date": cert.get("issue_date", ""),
            "expiration_date": cert.get("expiration_date", ""),
            "credential_id": cert.get("credential_id", ""),
            "credential_url": cert.get("credential_url", "")
        })

    return ResumeContentSchema(
        personal_info=personal_info,
        summary=summary,
        experience=experiences,
        education=education,
        skills=profile_dict.get("skills", ["Python", "FastAPI", "SQL", "Git"]),
        projects=projects,
        certifications=certifications
    )


async def improve_bullet_with_ai(section: str, original_text: str, instruction: str, user_id: Optional[str] = None) -> Dict[str, str]:
    """Improve specific section text or bullet point using Hugging Face Qwen with input sanitization."""
    from app.integrations.huggingface.client import improve_bullet_with_hf
    return await improve_bullet_with_hf(section=section, original_text=original_text, instruction=instruction, user_id=user_id)


async def save_generated_resume(db: AsyncSession, user: User, title: str, profile_snapshot: Dict[str, Any], content: ResumeContentSchema) -> Resume:
    """Save generated resume into DB and initialize Version 1."""
    content_dict = content.model_dump()
    resume = Resume(
        user_id=user.id,
        title=title,
        version=1,
        profile_snapshot=profile_snapshot,
        content=content_dict,
    )
    db.add(resume)
    await db.flush()

    v1 = ResumeVersion(
        resume_id=resume.id,
        version_number=1,
        title=title,
        content=content_dict,
        change_summary="Initial generated version",
    )
    db.add(v1)

    await db.commit()
    await db.refresh(resume)
    return resume


async def get_user_resumes(db: AsyncSession, user_id: str) -> list[Resume]:
    """Fetch all resumes for a user."""
    stmt = select(Resume).where(Resume.user_id == user_id).order_by(Resume.updated_at.desc())
    result = await db.execute(stmt)
    return list(result.scalars().all())


async def get_resume_by_id(db: AsyncSession, resume_id: str, user_id: str) -> Optional[Resume]:
    """Fetch single resume by ID for user."""
    stmt = select(Resume).where(Resume.id == resume_id, Resume.user_id == user_id)
    result = await db.execute(stmt)
    return result.scalar_one_or_none()

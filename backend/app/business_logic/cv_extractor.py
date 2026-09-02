import io
import json
import logging
import re
from typing import Any, Dict, List, Optional

from fastapi import HTTPException, status
from pypdf import PdfReader
import docx

logger = logging.getLogger(__name__)


# =============================================================================
# STEP 1: Extract raw text from the uploaded file
# =============================================================================

def clean_extracted_text(text: str) -> str:
    """Normalize extracted text: fix bullets, collapse whitespace/control chars."""
    if not text:
        return ""
    text = re.sub(r"[•▪►●◆★✓✔■–—]", "-", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def convert_document_to_markdown(file_bytes: bytes, filename: str) -> str:
    """
    Convert an uploaded PDF or DOCX into plain, lightly-structured Markdown.
    We deliberately do NOT try to guess resume sections/fields here — that job
    belongs entirely to the LLM in step 2. This function's only job is to turn
    the file into clean readable text with headings/bullets/tables preserved.
    """
    lower_filename = filename.lower()
    lines: List[str] = []

    if lower_filename.endswith(".pdf"):
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    lines.extend(page_text.splitlines())
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Unable to read PDF file. Please ensure it is a valid, unencrypted PDF.",
            )

    elif lower_filename.endswith(".docx"):
        try:
            doc = docx.Document(io.BytesIO(file_bytes))

            for para in doc.paragraphs:
                p_text = para.text.strip()
                if not p_text:
                    continue
                style_name = para.style.name.lower() if para.style else ""
                if "heading 1" in style_name:
                    lines.append(f"# {p_text}")
                elif "heading 2" in style_name or "heading 3" in style_name:
                    lines.append(f"## {p_text}")
                elif "bullet" in style_name or "list" in style_name:
                    lines.append(f"- {p_text}")
                else:
                    lines.append(p_text)

            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        lines.append(" | ".join(cells))

        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Unable to read DOCX document. Please ensure it is a valid Word (.docx) document.",
            )

    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Unsupported file format. Only PDF (.pdf) and Word (.docx) documents are allowed.",
        )

    cleaned_lines = [clean_extracted_text(line) for line in lines if line.strip()]
    md_document = "\n".join(cleaned_lines).strip()

    if not md_document or len(md_document) < 10:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="The uploaded file contains no extractable text. Please ensure it is a valid document with selectable text.",
        )

    return md_document


def extract_raw_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Extract the document as clean Markdown text (step 1 of the pipeline)."""
    return convert_document_to_markdown(file_bytes, filename)


# =============================================================================
# STEP 2: The fixed JSON schema we always want back, no matter what
# =============================================================================

def empty_cv_skeleton() -> Dict[str, Any]:
    """The exact shape the rest of the app expects. Used as defaults/fallback."""
    return {
        "personal_info": {
            "full_name": "",
            "professional_title": "",
            "email": "",
            "phone": "",
            "location": "",
            "linkedin_url": "",
            "github_url": "",
            "portfolio_url": "",
            "summary": "",
        },
        "experience": [],
        "education": [],
        "skills": [],
        "projects": [],
        "certifications": [],
    }


# The actual extraction prompt lives in extract_cv_with_hf() in the HF client —
# that's the single source of truth for what we ask the model to do, so we
# don't duplicate a second prompt here.


# =============================================================================
# STEP 3: Robustly parse + normalize whatever the LLM sends back
# =============================================================================

def extract_json_from_llm_output(text: Any) -> Optional[dict]:
    """
    LLMs often wrap JSON in ```json fences or add stray text around it.
    This pulls out a usable dict no matter how the response is formatted.
    """
    if isinstance(text, dict):
        return text
    if not isinstance(text, str) or not text.strip():
        return None

    cleaned = text.strip()
    cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", cleaned, flags=re.MULTILINE).strip()

    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass

    # Last resort: grab the first {...} block in the text
    match = re.search(r"\{[\s\S]*\}", cleaned)
    if match:
        try:
            return json.loads(match.group(0))
        except json.JSONDecodeError:
            return None

    return None


def normalize_cv_data(data: Optional[dict]) -> Dict[str, Any]:
    """
    Guarantees the output always has exactly the fields the app expects,
    with correct types. This is what actually fixes 'missing/wrong fields' —
    no matter what the LLM sends back, the caller always gets a complete,
    predictable object.
    """
    result = empty_cv_skeleton()
    if not isinstance(data, dict):
        return result

    personal = data.get("personal_info")
    if isinstance(personal, dict):
        for key in result["personal_info"]:
            value = personal.get(key)
            result["personal_info"][key] = value if isinstance(value, str) else result["personal_info"][key]

    for list_field in ("experience", "education", "skills", "projects", "certifications"):
        value = data.get(list_field)
        if isinstance(value, list):
            result[list_field] = value

    return result


# =============================================================================
# STEP 4: The actual pipeline — call the LLM, normalize, return
# =============================================================================

async def parse_cv_text_with_llm(raw_text: str, job_title: str = "", user_id: Optional[str] = None) -> Dict[str, Any]:
    """
    AI-first CV extraction pipeline:
      1. Send raw markdown CV directly to Hugging Face Qwen model (extract_cv_with_hf).
      2. Model performs deep extraction across Personal Info, Experience, Education, Skills, Projects, and Certifications.
      3. Normalize the AI output to the guaranteed schema and return.
    """
    from app.integrations.huggingface.client import extract_cv_with_hf, get_hf_key_manager

    km = get_hf_key_manager()
    if not km.has_keys():
        logger.warning("==> [CV Parser] No HF API keys configured. Returning empty skeleton.")
        return empty_cv_skeleton()

    try:
        logger.info(f"==> [CV Parser] Calling Hugging Face AI model for extraction ({len(raw_text)} chars)...")
        ai_result = await extract_cv_with_hf(raw_text, job_title=job_title, user_id=user_id)
    except Exception as e:
        logger.error(f"==> [CV Parser] AI extraction raised an exception: {e}", exc_info=True)
        ai_result = None

    if not ai_result:
        logger.warning("==> [CV Parser] AI model returned no parsed data.")
        return empty_cv_skeleton()

    parsed = extract_json_from_llm_output(ai_result)
    normalized = normalize_cv_data(parsed)

    logger.info(
        f"==> [CV Parser] AI extraction complete: "
        f"{len(normalized['experience'])} experiences, "
        f"{len(normalized['education'])} education, "
        f"{len(normalized['skills'])} skills, "
        f"{len(normalized['projects'])} projects."
    )
    return normalized


# =============================================================================
# STEP 5: Skill Categorization Helper (used by Resume Generator / Template)
# =============================================================================

SOFT_SKILLS_KEYWORDS = {
    "communication", "interpersonal", "teamwork", "collaboration", "organizational",
    "time management", "problem-solving", "attention to detail", "leadership",
    "project management", "critical thinking", "analytical", "creative thinking",
    "adaptability", "flexibility", "reliability", "accountability", "initiative",
    "presentation", "public speaking", "negotiation", "conflict resolution",
    "customer service", "client management", "stakeholder management", "ms office",
    "word", "excel", "powerpoint", "event coordination", "event management",
    "recruitment", "recruiting", "hiring", "onboarding", "training",
    "mentoring", "coaching", "employee relations", "hr", "human resources",
    "networking", "relationship building", "multitasking", "prioritization"
}


def categorize_skills(skills_list: List[str]) -> tuple[List[str], List[str]]:
    """Separate technical skills from soft/professional skills for resume template rendering."""
    technical_skills = []
    soft_skills = []

    for skill in skills_list:
        skill_lower = skill.lower().strip()
        is_soft_skill = False
        for soft_keyword in SOFT_SKILLS_KEYWORDS:
            if soft_keyword in skill_lower or skill_lower in soft_keyword:
                is_soft_skill = True
                break

        if is_soft_skill:
            soft_skills.append(skill)
        else:
            technical_skills.append(skill)

    return technical_skills, soft_skills